import json
import uuid
import openai
from docx import Document

def prepare(filename):
    target_file = "tr-{}.txt".format(uuid.uuid4())
    
    f = open(filename)
    data = json.load(f)
    x = open(target_file, "a")
    for i in data['results']['audio_segments']:
        x.write("{}: {}".format(i['speaker_label'], i['transcript']))
        x.write("\n")

    x.close()
    f.close()

    return target_file

def generate_meeting_minute(summary_file, openai_key):
    openai.api_key = openai_key
    with open(summary_file, 'r') as file:
        content = file.read()

    messages = [
        {"role": "system", "content": "You are a professional assistant that create meeting minutes."},
        {"role": "user", "content": f"Please summary this conversation record to meeting minute format. Define the date based on today, participant, agenda, key points discussed, action items, and conclussion. \n\n{content}"}
    ]

    # Generate a response from ChatGPT using GPT-4
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=messages
    )
    return response.choices[0].message['content'].strip()

def save_to_docx(content):
    output_file_path = "meeting-minute_{}.docx".format(uuid.uuid4())
    # Create a Document object
    doc = Document()
    
    # Add a heading
    doc.add_heading('Meeting Minutes', level=1)
    
    # Add the content to the document
    doc.add_paragraph(content)
    
    # Save the document
    doc.save(output_file_path)
    print(f"Meeting minutes saved to {output_file_path}")